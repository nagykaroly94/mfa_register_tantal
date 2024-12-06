import qrcode
from PIL import Image
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import tkinter as tk
from tkinter import ttk
from tkinter import messagebox

def generate_qr_code(text, output_path):
    """QR-kód generálása és mentése."""
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10,
        border=4,
    )
    qr.add_data(text)
    qr.make(fit=True)
    img = qr.make_image(fill="black", back_color="white")
    img.save(output_path)

def create_document_from_treeview(treeview):
    """Kiválasztott Treeview-sor alapján dokumentum létrehozása."""
    selected_item = treeview.selection()
    if not selected_item:
        messagebox.showerror("Hiba", "Kérjük, válasszon ki egy sort!")
        return

    # Az adatok kinyerése a kijelölt sorból
    selected_data = treeview.item(selected_item[0])['values']
    nev = selected_data[1]  # Feltételezzük, hogy a név a második oszlopban van
    helyreallitasi_kod = selected_data[2]  # Helyreállítási kód például a harmadik oszlopban
    qr_kod_szoveg = f"{nev} - {helyreallitasi_kod}"  # QR-kód szövege

    # QR-kód generálása
    qr_kod_path = "generated_qr_code.png"
    generate_qr_code(qr_kod_szoveg, qr_kod_path)

    # Dokumentum létrehozása
    doc = Document()

    # Cím
    doc.add_heading("Kétlépcsős Hitelesítés Beállítási Útmutató", level=1)

    # Bevezető szöveg
    doc.add_paragraph(
        "Ez a dokumentum tartalmazza a kétlépcsős hitelesítés beállításához szükséges információkat. "
        "Kérjük, kövesse az alábbi lépéseket az autentikátor alkalmazás beállításához, és őrizze meg ezt a dokumentumot biztonságos helyen."
    )

    # Beállítási lépések
    doc.add_heading("Beállítási lépések", level=2)
    steps = [
        "Töltse le és telepítse az Google Authenticator alkalmazást a mobiltelefonjára Google Play Áruházból vagy az App Store-ból.",
        "Nyissa meg az alkalmazást, és válassza az „Új fiók hozzáadása” vagy „QR-kód beolvasása” opciót.",
        "Olvassa be az alábbi QR-kódot az alkalmazás segítségével.",
        "Az alkalmazás generál egy hatjegyű kódot, amelyet használhat a bejelentkezések során.",
        "Jegyezze fel és tárolja biztonságosan a törlőkódokat (helyreállítási kódokat), amelyek az alkalmazás elvesztése esetén használhatók."
    ]
    for step in steps:
        doc.add_paragraph(step, style="List Number")

    # QR-kód beszúrása
    doc.add_heading("QR-kód", level=2)
    doc.add_paragraph("Az alábbi QR-kódot olvassa be az autentikátor alkalmazásával:")

    doc.add_picture(qr_kod_path, width=Pt(300))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Helyreállítási kód
    doc.add_heading("Helyreállítási kód", level=2)
    p = doc.add_paragraph()
    run = p.add_run(helyreallitasi_kod)
    run.font.size = Pt(18)
    run.bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Figyelmeztetés
    doc.add_paragraph(
        "FONTOS: Ezt a dokumentumot tárolja biztonságos helyen! Ne ossza meg a QR-kódot vagy a helyreállítási kódot senkivel."
    )

    # Dokumentum mentése
    file_name = f"{nev}_ketlepcsos_hitelesites.docx"
    doc.save(file_name)
    print(f"A dokumentum elkészült: {file_name}")
    messagebox.showinfo("Siker", f"A dokumentum mentve: {file_name}")

# Példa a Treeview és függvény használatára
def example_treeview_selection():
    root = tk.Tk()
    root.title("Példa Treeview")

    # Treeview létrehozása
    treeview = tk.ttk.Treeview(root, columns=("ID", "Név", "Helyreállítási kód"), show="headings")
    treeview.heading("ID", text="ID")
    treeview.heading("Név", text="Név")
    treeview.heading("Helyreállítási kód", text="Helyreállítási kód")
    
    # Adatok beszúrása
    treeview.insert("", tk.END, values=(1, "Kiss Péter", "12345678"))
    treeview.insert("", tk.END, values=(2, "Nagy Anna", "87654321"))

    treeview.pack(fill=tk.BOTH, expand=True)

    # Gomb a dokumentum generálásához
    generate_button = tk.Button(root, text="Dokumentum generálása", command=lambda: create_document_from_treeview(treeview))
    generate_button.pack(pady=10)

    root.mainloop()

example_treeview_selection()
