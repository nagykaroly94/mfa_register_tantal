import mysql.connector
from mysql.connector import Error
import configparser
import tkinter as tk
from tkinter import filedialog
from tkinter import ttk
import tkinter.messagebox
from tkinter import messagebox
import qrcode
from PIL import ImageTk
from PIL import ImageGrab
import cv2
import numpy as np
import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor

kivalasztott_png = None
value = None
search_entry = None

def db_kapcsolodas():
    config = configparser.ConfigParser()
    config.read("config.ini")
    try:
        connection = mysql.connector.connect(
            host=config["database"]["host"],
            user=config["database"] ["user"],
            password=config["database"]["password"],
            database=config["database"]["database"]
        )
        return connection
    except Error as e:
        print(f"Hiba történt: {e}")
        return None

def png_kivalasztasa():
    global kivalasztott_png
    kivalasztott_png = filedialog.askopenfilename(
        title="Válassz egy QR kódot",
        filetypes=[('image files', '.png;.jpg')]
    )
    if kivalasztott_png:
        qrkod_png.set(kivalasztott_png)
        print(kivalasztott_png)
        qr_read()

def qr_read():
    global value
    if not kivalasztott_png:
        msg1 = ("Nem választottál ki fájlt!")
        tkinter.messagebox.showerror(title="Hiba", message=msg1)
        return

    if not os.path.exists(kivalasztott_png):
        msg2 = (f"A fájl nem található: {kivalasztott_png}")
        tkinter.messagebox.showerror(title="Hiba", message=msg2)
        return

    img = cv2.imdecode(np.fromfile(kivalasztott_png, dtype=np.uint8), cv2.IMREAD_COLOR)
    if img is None:
        msg3 = ("Nem sikerült betölteni a képet!")
        tkinter.messagebox.showerror(title="Hiba", message=msg3)
    else:
        # QR kód olvasása
        detector = cv2.QRCodeDetector()
        
        # A detectAndDecode() metódus használata a QR kód dekódolásához
        value, pts, qr_code = detector.detectAndDecode(img)

        if value:
            print("QR kód szövege:", value)
        else:
            msg4 = ("Nem találtunk QR kódot.")
            tkinter.messagebox.showerror(title="Hiba", message=msg4)

def fetch_data():
    """Adatok lekérdezése a MySQL adatbázisból."""
    try:
        # Hívjuk a db_kapcsolodas() függvényt, hogy kapcsolatot létesítsünk
        connection = db_kapcsolodas()

        # Ellenőrizzük, hogy a kapcsolat sikerült-e
        if connection is None:
            print("Nem sikerült kapcsolatot létesíteni az adatbázissal.")
            return []

        cursor = connection.cursor()
        cursor.execute("SELECT id, nev, cegnev, felhasznalonev, jelszo, telefonszam, emailcim, torlokod FROM adatok")
        rows = cursor.fetchall()
        
        for row in treeview.get_children():
            treeview.delete(row)
        
        # Új sorok hozzáadása a Treeview-hez
        for row in rows:
            treeview.insert("", tk.END, values=row)        
        return rows

    except mysql.connector.Error as e:
        print(f"Hiba történt az adatbázis lekérdezés során: {e}")
        return []
    finally:
        # Zárjuk le a kapcsolatot, ha van
        if connection and connection.is_connected():
            connection.close()

def default_data():
    search_entry.delete(0, tk.END)

def default_search():
    default_data()
    fetch_data()

def copy_to_clipboard():
    """Az aktuálisan kijelölt cella másolása a vágólapra."""
    selected_item = treeview.focus()  # Az aktuálisan kijelölt sor
    if not selected_item:
        return  # Ha nincs sor kijelölve, kilép

    # Az aktuálisan kijelölt oszlop indexének meghatározása a treeview.selection() alapján
    column_index = treeview.identify_column(treeview.winfo_pointerx() - treeview.winfo_rootx())  # Az oszlop azonosítása
    column_index = int(column_index.replace("#", "")) - 1  # Az oszlop indexe
    
    # Ha nem az aktuálisan kijelölt oszlopot próbáljuk másolni
    if column_index < 0:
        return
    
    cell_value = treeview.item(selected_item)["values"][column_index]  # Az adott cella értéke
    root.clipboard_clear()
    root.clipboard_append(cell_value)
    root.update()  # A vágólap frissítése
    #print(f"Másolt adat: {cell_value} (Oszlop: {columns[column_index]})")
    label_selected.config(text=f"Másolt adat: {cell_value} (Oszlop: {columns[column_index]})")

def show_selected_column(event):
    """Az aktuálisan kijelölt oszlop és sor megjelenítése, és az Entry mezők frissítése."""
    selected_item = treeview.focus()
    if not selected_item:
        label_selected.config(text="Nincs kijelölés")
        return

    # Az aktuális oszlop meghatározása
    column = treeview.identify_column(event.x)
    column_index = int(column.replace("#", "")) - 1  # Oszlop indexe (0-alapú)
    
    # A kijelölt sor értékei
    values = treeview.item(selected_item)["values"]

    # Ellenőrizzük, hogy van-e adat az adott oszlopban
    if column_index >= len(values):
        label_selected.config(text="Hiányzó adat az oszlopban")
        return

    # Aktuális cella értéke
    cell_value = values[column_index]
    label_selected.config(text=f"Kijelölt oszlop: {columns[column_index]}, Adat: {cell_value}")

    # Az összes adat betöltése az Entry mezőkbe
    if len(values) >= 7:
        nev_entry.delete(0, tk.END)
        nev_entry.insert(0, values[1]) 

        cegnev_entry.delete(0, tk.END)
        cegnev_entry.insert(0, values[2]) 

        felhasznalonev_entry.delete(0, tk.END)
        felhasznalonev_entry.insert(0, values[3])

        jelszo_entry.delete(0, tk.END)
        jelszo_entry.insert(0, values[4])

        telefon_entry.delete(0, tk.END)
        telefon_entry.insert(0, values[5])

        email_entry.delete(0, tk.END)
        email_entry.insert(0, values[6])

        torlo_entry.delete(0, tk.END)
        torlo_entry.insert(0, values[7])
        add_button.config(state="disabled")

def generate_qr(data):
    """QR kód generálása az adott szövegből."""
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10,
        border=4,
    )
    qr.add_data(data)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    return img

def get_selected_item_id():
    """Lekéri a kiválasztott sor ID-ját a Treeview-ból."""
    selected_item = treeview.focus()  # Kiválasztott sor azonosítója
    if not selected_item:
        return None  # Ha nincs kijelölt sor, visszatérünk None-nal

    values = treeview.item(selected_item, "values")  # A kiválasztott sor értékei
    item_id = values[0]  # Az ID az első oszlopban (index 0)
    
    return item_id

def get_qr_code_from_db(item_id):
    """Lekéri a QR kódot az adatbázisból a kiválasztott sor ID-ja alapján."""
    try:
        connection = db_kapcsolodas()
        if connection is None:
            tkinter.messagebox.showerror(title="Kapcsolódási hiba", message="Nem sikerült kapcsolatot létesíteni az adatbázissal.")
            return None
        cursor = connection.cursor()
        cursor.execute("SELECT `qr_kod` FROM `adatok` WHERE `ID` = %s", (item_id,))
        qr_code = cursor.fetchone()

        if qr_code:
            return qr_code[0]  # A QR kód értéke az első oszlopban
        else:
            tkinter.messagebox.showwarning(title="Hiba", message=msg)
            msg("Nem található QR kód ezzel az ID-val!")
            return None

    except mysql.connector.Error as e:
        tkinter.messagebox.showerror(title="Kapcsolódási hiba", message=msg)
        msg=("Hiba történt az adatbázis lekérdezésekor:", e)
        return None

    finally:
        if connection and connection.is_connected():
            connection.close()

def on_tree_select(event):
    item_id = get_selected_item_id()
    if item_id:
        qr_code = get_qr_code_from_db(item_id)
        if qr_code is None:
             tkinter.messagebox.showwarning(title="Hiba", message="Nem találtunk QR kódot ehhez az ID-hoz.")  # Hibaüzenet ha nincs QR kód
             return
    """A kijelölt sor QR kód oszlopából QR kód generálása és megjelenítése."""
    selected_item = treeview.focus()  # Kijelölt sor azonosítója
    if not selected_item:
        qr_text.set("Válassz egy sort a QR kód megjelenítéséhez!")
        qr_canvas.delete("all")  # Canvas törlése
        return

    qr_data = qr_code
    qr_text.set(qr_data if qr_data else "Nincs megjeleníthető QR kód!")

    if qr_data:
        img = generate_qr(qr_data)
        img = img.resize((250, 250))  # Méret igazítása
        tk_image = ImageTk.PhotoImage(img)

        qr_canvas.delete("all")
        
        # QR kód középre helyezése
        qr_canvas.update()
        canvas_width = qr_canvas.winfo_width()
        canvas_height = qr_canvas.winfo_height()
        x_center = (canvas_width - img.width) // 2
        y_center = (canvas_height - img.height) // 2
        
        qr_canvas.create_image(x_center, y_center, anchor=tk.NW, image=tk_image)
        qr_canvas.image = tk_image

def adjust_column_width(tree, columns, data):
    """Automatikusan beállítja az oszlopok szélességét a legszélesebb tartalom alapján."""
    for col_index, col_name in enumerate(columns):
        # Legszélesebb adat hosszának meghatározása
        max_width = max(len(str(row[col_index])) for row in data)  
        
        # A betűmérethez igazított szélesség kiszámítása
        font_size_factor = 14  # Alap szorzó, finomhangolható
        column_width = max(max_width * font_size_factor, 100)  # Minimalizált szélesség biztosítása
        
        # Oszlop szélességének beállítása
        tree.column(col_name, width=column_width)

def kijelolt_adat_torles():
    # Kiválasztott sor lekérdezése
    selected_item = treeview.selection()

    if selected_item:
        # Az első oszlop (azonosito) értékének lekérdezése
        selected_data = treeview.item(selected_item[0])['values'][0]
        selected_username = treeview.item(selected_item[0])['values'][1]
        
        # Első megerősítő párbeszédablak
        confirm = messagebox.askyesno("Megerősítés", f"Biztosan törölni szeretnéd a(z) {selected_username} ügyfelet?")
        
        if confirm:
            # Második megerősítő párbeszédablak
            confirm_final = messagebox.askyesno("Végső megerősítés", f"Ez a művelet nem vonható vissza! Biztos vagy benne, hogy törlöd a(z) {selected_username} ügyfelet?")
            
            if confirm_final:
                # Törlés az adatbázisból
                mydb = db_kapcsolodas()
                if mydb is None:
                    messagebox.showerror("Hiba", "Nem sikerült csatlakozni az adatbázishoz.")
                    return
                
                try:
                    mycursor = mydb.cursor()
                    sql = "DELETE FROM adatok WHERE id = %s"
                    val = (selected_data,)
                    mycursor.execute(sql, val)
                    mydb.commit()  # Az adatbázis módosításainak végrehajtása
                    messagebox.showinfo("Sikeres törlés", f"{selected_username} ügyfél törölve.")
                    
                    # Frissítsd a treeview-t
                    treeview.delete(selected_item[0])

                    # Új adatok betöltése a fetch_data() meghívásával
                    fetch_data()
                    
                except mysql.connector.Error as e:
                    mydb.rollback()  # Ha hiba történik, vonja vissza a módosítást
                    messagebox.showerror("Hiba", f"Hiba történt a törlés során: {e}")
                finally:
                    mydb.close()
                
            else:
                messagebox.showwarning("Figyelmeztetés", "A végső megerősítés hiányában a törlési művelet megszakítva.")
        else:
            messagebox.showwarning("Figyelmeztetés", "A törlés művelet megszakítva.")
    else:
        messagebox.showerror("Hiba", "Nincs kijelölve sor a törléshez.")

def ugyfel_megad():
    global value
    nev = nev_entry.get()
    cegnev = cegnev_entry.get()
    felhasznalonev = felhasznalonev_entry.get()
    jelszo = jelszo_entry.get()
    telefon = telefon_entry.get()
    email = email_entry.get()
    torlo = torlo_entry.get()
    qrkod = value


    if not (nev and cegnev and felhasznalonev and jelszo and telefon and email and torlo):
        messagebox.showerror("Hiba", "Minden mezőt ki kell tölteni a sikeres hozzáadáshoz!")
        return
    if qrkod == None :
        messagebox.showerror("Hiba", "Szükséges egy érvényes QR kódot is kiválasztani!")
        return
    mydb = db_kapcsolodas()

    if mydb is None:
        messagebox.showerror("Hiba", "Nem sikerült csatlakozni az adatbázishoz.")
        return
        
    mycursor = mydb.cursor()

    sql= """
        INSERT INTO adatok (nev, cegnev, felhasznalonev, jelszo, telefonszam, emailcim, torlokod, qr_kod) 
        VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
        """
    values = (nev, cegnev, felhasznalonev, jelszo, telefon, email, torlo, qrkod)
    try:
            # SQL lekérdezés végrehajtása
        mycursor.execute(sql, values)
        mydb.commit()
        fetch_data()
        messagebox.showinfo("Siker", "Az Adatok sikeresen rögzítve az adatbázisban.")
    except Error as e:
        messagebox.showerror("Hiba", f"Hiba történt: {e}")
    finally:
            # Kapcsolat lezárása
        mydb.close()
        nev_entry.delete(0, tk.END)
        cegnev_entry.delete(0, tk.END)
        felhasznalonev_entry.delete(0, tk.END)
        jelszo_entry.delete(0, tk.END)
        telefon_entry.delete(0, tk.END)
        email_entry.delete(0, tk.END)
        torlo_entry.delete(0, tk.END)
        qrkod_png.set("C:/")

def search_ceg():
    global search_entry
    # Keresési kulcsszó lekérdezése a beviteli mezőből
    search_term = search_entry.get()
    
    # MySQL adatbázishoz való csatlakozás
    mydb = db_kapcsolodas()
    if mydb is None:
        messagebox.showerror("Hiba", "Nem sikerült csatlakozni az adatbázishoz.")
        return
    mycursor = mydb.cursor()
    
    # SQL lekérdezés keresési feltétellel az "azonosito" oszlopra
    sql = "SELECT id, nev, cegnev, felhasznalonev, jelszo, telefonszam, emailcim, torlokod FROM adatok WHERE cegnev LIKE %s ORDER BY id ASC"
    search_pattern = f"%{search_term}%"
    mycursor.execute(sql, (search_pattern,))  # A % jelek a keresési mintát határozzák meg
    
    sorok = mycursor.fetchall()
    
    # Töröljük a Treeview régi adatait
    for item in treeview.get_children():
        treeview.delete(item)
    
    # Új adatok hozzáadása a Treeview-hez
    for sor in sorok:
        print(sor)
        treeview.insert("", tk.END, values=sor)
        
    mydb.close()

def ugyfel_frissites():
    global ertek_mod
    # Kiválasztott sor lekérdezése
    selected_item = treeview.selection()

    if selected_item:
        # Az első oszlop (azonosito) értékének lekérdezése
        selected_data = treeview.item(selected_item[0])['values'][0]
        selected_username = treeview.item(selected_item[0])['values'][1]

        # Első megerősítő párbeszédablak
        confirm = messagebox.askyesno("Megerősítés", f"Biztosan frissíteni szeretnéd a(z) {selected_username} felhasználót?")
        
        if confirm:
            # Második megerősítő párbeszédablak
            confirm_final = messagebox.askyesno("Végső megerősítés", f"Ez a művelet nem vonható vissza! Biztos vagy benne, hogy frissíted a(z) {selected_username} felhasználót?")
            
            if confirm_final:
                # Új értékek lekérdezése az entry mezőkből
                uj_nev = nev_entry.get()
                uj_cegnev = cegnev_entry.get()
                uj_felhasznalonev = felhasznalonev_entry.get()
                uj_jelszo = jelszo_entry.get()
                uj_telefon = telefon_entry.get()
                uj_email = email_entry.get()
                uj_torlo = torlo_entry.get()
                uj_qrkod = value
                
                # Ellenőrizzük, hogy vannak-e kitöltött értékek
                if not (uj_nev and uj_cegnev and uj_felhasznalonev and uj_jelszo and uj_telefon and uj_email and uj_torlo):
                    messagebox.showerror("Hiba", "A Sikeres frissítéshez minden mezőt ki kell tölteni!")
                    return
                if uj_qrkod == None:
                    messagebox.showerror("Hiba", "A Sikeres frissítéshez szükséges egy érvényes QR kódot is kiválasztani.")
                    return

                # Frissítés az adatbázisban
                try:
                    mydb = db_kapcsolodas()
                    if mydb is None:
                        messagebox.showerror("Hiba", "Nem sikerült csatlakozni az adatbázishoz.")
                        return

                    mycursor = mydb.cursor()
                    sql = "UPDATE adatok SET nev = %s, cegnev = %s, felhasznalonev = %s, jelszo = %s, telefonszam = %s, emailcim = %s, torlokod = %s, qr_kod = %s WHERE id = %s"
                    val = (uj_nev, uj_cegnev, uj_felhasznalonev, uj_jelszo, uj_telefon, uj_email, uj_torlo, uj_qrkod, selected_data)
                    mycursor.execute(sql, val)
                    mydb.commit()
                except Exception as e:
                    print("Hiba történt:", e)
                    messagebox.showerror("Hiba", f"Hiba történt: {e}")
                finally:
                    if mydb:
                        mydb.close()
                    fetch_data()
                    nev_entry.delete(0, tk.END)
                    cegnev_entry.delete(0, tk.END)
                    felhasznalonev_entry.delete(0, tk.END)
                    telefon_entry.delete(0, tk.END)
                    email_entry.delete(0, tk.END)
                    torlo_entry.delete(0, tk.END)
                    jelszo_entry.delete(0,tk.END)
                    qrkod_png.set("C:/")

                messagebox.showinfo("Sikeres frissítés", f"{selected_username} felhasználó adatai frissítve.")
            else:
                messagebox.showwarning("Figyelmeztetés", "A végső megerősítés hiányában a frissítési művelet megszakítva.")
        else:
            messagebox.showwarning("Figyelmeztetés", "A frissítési művelet megszakítva.")
    else:
        messagebox.showerror("Hiba", "Nincs kijelölve sor a frissítéshez.")

def export_document(treeview, qr_canvas):
    # Kiválasztott sor adatai
    selected_item = treeview.selection()
    if not selected_item:
        messagebox.showerror("Hiba", "Kérjük, válasszon ki egy sort!")
        return

    selected_data = treeview.item(selected_item[0])['values']
    nev = selected_data[1]  # Név a második oszlopban
    helyreallitasi_kod = selected_data[7]  # Törlő kód a hetedik oszlopban

    # QR-kód mentése a Canvasról
    qr_kod_path = f"{nev}_qr_kod.png"
    try:
        x0 = qr_canvas.winfo_rootx()
        y0 = qr_canvas.winfo_rooty()
        x1 = x0 + qr_canvas.winfo_width()
        y1 = y0 + qr_canvas.winfo_height()
        ImageGrab.grab(bbox=(x0, y0, x1, y1)).save(qr_kod_path)
    except Exception as e:
        messagebox.showerror("Hiba", f"Nem sikerült menteni a QR-kódot: {e}")
        return

    # Dokumentum létrehozása
    doc = Document()

    # Cím
    doc.add_heading("Kétlépcsős Hitelesítés Beállítási Útmutató", level=1)

    # Bevezető szöveg
    doc.add_paragraph(
        "Ez a dokumentum tartalmazza a kétlépcsős hitelesítés beállításához szükséges információkat. "
        "Kérjük, kövesse az alábbi lépéseket az autentikátor alkalmazás beállításához, és őrizze meg ezt a dokumentumot biztonságos helyen."
    )

    # Beállítási lépések
    doc.add_heading("Beállítási lépések", level=2)
    steps = [
        "Töltse le és telepítse az Google Authenticator alkalmazást a mobiltelefonjára Google Play Áruházból vagy az App Store-ból.",
        "Nyissa meg az alkalmazást, és válassza az „Új fiók hozzáadása” vagy „QR-kód beolvasása” opciót.",
        "Olvassa be az alábbi QR-kódot az alkalmazás segítségével.",
        "Az alkalmazás generál egy hatjegyű kódot, amelyet használhat a bejelentkezések során.",
        "Jegyezze fel és tárolja biztonságosan a törlőkódokat, amelyek az alkalmazás elvesztése esetén használhatók."
    ]
    for step in steps:
        doc.add_paragraph(step, style="List Number")

    # QR-kód rész
    doc.add_heading("QR-kód", level=2)
    doc.add_paragraph("Az alábbi QR-kódot olvassa be az autentikátor alkalmazásával:")
    doc.add_picture(qr_kod_path, width=Pt(500))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # QR-kód alatti szöveg vége az egyenlőségjel után
    qr_canvas_text = qr_text.get()
    if "=" in qr_canvas_text:
       qr_adat = qr_canvas_text.split("=")[-1].strip()  # Az egyenlőségjel utáni rész
    else:
        qr_adat = "Nincs egyenlőségjel a szövegben"



    # Másik adat hozzáadása a dokumentumhoz
    doc.add_heading("QR-kód nélkül", level=2)
    doc.add_paragraph("Ha a készüléke nem olvassa be a képet, a mobilalkalmazásban válassza a manuális beállítást és adja meg az alábbi karaktersort. A fiók nevét a kód beírásánál Ön határozhatja meg (célszerű beszédes elnevezést megadni, mint például az ügyfélkapus felhasználónév).")
    p = doc.add_paragraph()
    run = p.add_run(qr_adat)
    run.font.size = Pt(18)
    run.bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    # törlő kód
    doc.add_heading("Törlő kód", level=2)
    p = doc.add_paragraph()
    run = p.add_run(str(helyreallitasi_kod))
    run.font.size = Pt(18)
    run.bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Figyelmeztetés
    paragraph = doc.add_paragraph()
    important_run = paragraph.add_run("FONTOS: Ezt a dokumentumot tárolja biztonságos helyen! Ne ossza meg a QR-kódot vagy a törlő kódot senkivel.")
    important_run.font.size = Pt(16)
    important_run.font.color.rgb = RGBColor(255, 0, 0)
    important_run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Dokumentum mentése
    file_name = f"{nev}_ketlepcsos_hitelesites.docx"
    try:
        doc.save(file_name)
        messagebox.showinfo("Siker", f"A dokumentum mentve: {file_name}")
        os.startfile(file_name)
    except Exception as e:
        messagebox.showerror("Hiba", f"Nem sikerült menteni a dokumentumot: {e}")
    finally:
        # Tisztítsd meg az ideiglenes QR-kód fájlt
        if os.path.exists(qr_kod_path):
            os.remove(qr_kod_path)

def empty_entry():
    nev_entry.delete(0, tk.END)
    cegnev_entry.delete(0, tk.END)
    felhasznalonev_entry.delete(0, tk.END)
    jelszo_entry.delete(0, tk.END)
    telefon_entry.delete(0, tk.END)
    email_entry.delete(0, tk.END)
    torlo_entry.delete(0, tk.END)
    add_button.config(state="normal")
    qrkod_png.set("C:/")

# Főablak alapvető beállításai
img=("ikon.ico")
root = tk.Tk()
root.title("Tantál")
root.iconbitmap(img)
root.grid_columnconfigure(0, weight=1)
root.grid_rowconfigure(0, weight=1)
root.resizable(False, False)

# Rácshoz tartozó beállítások, frame-ek beálíltása
main_frame = tk.Frame(root)
main_frame.grid(row=0, column=0, sticky="nsew", columnspan=4)
root.rowconfigure(0, weight=1)
root.columnconfigure(0, weight=1)
main_frame.columnconfigure(0, weight=3)
main_frame.columnconfigure(1, weight=1)
main_frame.columnconfigure(2, weight=2)
main_frame.rowconfigure(0, weight=1)
main_frame.rowconfigure(2, weight=1)
treeview_frame = tk.Frame(main_frame)
treeview_frame.grid(row=0, column=0, sticky="nsew", columnspan=4)
qr_frame = tk.Frame(main_frame, padx=20, pady=15, bg="white")
qr_frame.grid(row=3, column=2, sticky="nsew", padx=15)
qr_frame.grid_columnconfigure(0, weight=1)
qr_frame.grid_rowconfigure(0, weight=1)
search_frame = tk.Frame(main_frame)
search_frame.grid(row=2, column=0, sticky="NSEW")
data_frame = tk.Frame(main_frame)
data_frame.grid(row=3, column=0, columnspan=2)

#Adatok megadása és funkció gombok létrehozása
label = tk.Label(data_frame, text="Név:", font=("Arial", 16))
label.grid(row=0, column=0, sticky="W", padx=10, pady=10)

nev_entry = tk.Entry(data_frame, font=("Arial", 16))
nev_entry.grid(row=0, column=1, sticky="W", padx=10, pady=10)

label = tk.Label(data_frame, text="Cégnév:", font=("Arial", 16))
label.grid(row=2, column=0, sticky="W", padx=10, pady=10)

cegnev_entry = tk.Entry(data_frame, font=("Arial", 16))
cegnev_entry.grid(row=2, column=1, padx=10, pady=10, sticky="w")

label = tk.Label(data_frame, text="Felhasználónév:", font=("Arial", 16))
label.grid(row=4, column=0, sticky="W", padx=10, pady=10)

felhasznalonev_entry = tk.Entry(data_frame, font=("Arial", 16))
felhasznalonev_entry.grid(row=4, column=1, padx=10, pady=10, sticky="w")

label = tk.Label(data_frame, text="Jelszó:", font=("Arial", 16))
label.grid(row=6, column=0, sticky="W", padx=10, pady=10)

jelszo_entry = tk.Entry(data_frame, font=("Arial", 16))
jelszo_entry.grid(row=6, column=1, padx=10, pady=10, sticky="w")

label = tk.Label(data_frame, text="Telefonszám:", font=("Arial", 16))
label.grid(row=8, column=0, sticky="W", padx=10, pady=10)

telefon_entry = tk.Entry(data_frame, font=("Arial", 16))
telefon_entry.grid(row=8, column=1, padx=10, pady=10, sticky="w")

label = tk.Label(data_frame, text="E-mail cím:", font=("Arial", 16))
label.grid(row=10, column=0, sticky="W", padx=10, pady=10)

email_entry = tk.Entry(data_frame, font=("Arial", 16))
email_entry.grid(row=10, column=1, padx=10, pady=10, sticky="w")

label = tk.Label(data_frame, text="QR kód:", font=("Arial", 16))
label.grid(row=14, column=0, sticky="W", padx=10, pady=10)

qrkod_png=tk.StringVar()
qrkod_png.set("C:/")
qrkod_entry = tk.Entry(data_frame, font=("Arial", 16),textvariable=qrkod_png)
qrkod_entry.grid(row=14, column=1, padx=10, pady=10, sticky="nsew", columnspan=4)

talloz_button = tk.Button(data_frame, font=("Arial", 16), text="QR kód kiválasztása", command=png_kivalasztasa)
talloz_button.grid(row=12, column=3, padx=10, pady=10, sticky="nsew", columnspan=2)

label = tk.Label(data_frame, text="Törlő kód:", font=("Arial", 16))
label.grid(row=12, column=0, sticky="W", padx=10, pady=10)

torlo_entry = tk.Entry(data_frame, font=("Arial", 16))
torlo_entry.grid(row=12, column=1, padx=10, pady=10)

add_button = tk.Button(data_frame, font=("Arial", 16), text="Új Ügyfél hozzáadása", command=ugyfel_megad)
add_button.grid(row=4, column=3, padx=10, pady=10, columnspan=2, sticky="nsew")

del_button = tk.Button(data_frame, font=("Arial", 16), text="Ügyfél törlése a rendszerből", command=kijelolt_adat_torles)
del_button.grid(row=6, column=3, padx=10, pady=10, columnspan=2, sticky="nsew")

mod_button = tk.Button(data_frame, font=("Arial", 16), text="Ügyfél adat módosítás", command=ugyfel_frissites)
mod_button.grid(row=8, column=3, padx=10, pady=10, columnspan=2, sticky="nsew")

ex_button = tk.Button(data_frame, font=("Arial", 16), text="Dokumentum létrehozása", command=lambda:export_document(treeview, qr_canvas))
ex_button.grid(row=10, column=3, padx=10, pady=10, columnspan=2, sticky="nsew")

# Treeview létrehozása
style = ttk.Style()
style.configure("Treeview", font=("Arial", 16))
style.configure("Treeview.Heading", font=("Arial", 12))  # Betűméret növelése
style.configure("Treeview", rowheight=30)  # Sorok (és fejlécek) magasságának növelése
columns = ["ID", "Név", "Cégnév vagy Adószám", "Felhasználónév", "Jelszó", "Telefonszám", "E-mail cím", "Törlőkód"]
treeview = ttk.Treeview(treeview_frame, columns=columns, show="headings")
treeview.grid(row=0, column=0, sticky="nsew")

# Oszlopok széllességének a beállítása
data = fetch_data()
adjust_column_width(treeview, columns, data)
for col in columns:
    treeview.heading(col, text=col)

# Görgetősáv hozzáadása
scrollbar = ttk.Scrollbar(treeview_frame, orient=tk.VERTICAL, command=treeview.yview)
treeview.configure(yscroll=scrollbar.set)
scrollbar.grid(row=0, column=2, sticky="ns")

# Label a kijelölt adat megjelenítéséhez
label_selected = tk.Label(main_frame, text="Nincs kijelölés", anchor="w", font=("Arial", 20))
label_selected.grid(row=1, column=0, sticky="ew")

#Kereséshez szükséges elemek
label_search = tk.Label(search_frame, text="Cég szerinti keresés:", font=("Arial", 20))
label_search.grid(row=0, column=0, sticky="NSEW")

search_entry = tk.Entry(search_frame, font=("Arial", 16))
search_entry.grid(row=0, column=1, padx=10, pady=10, sticky="nsew")

search_button = tk.Button(search_frame, font=("Arial", 16), text="Keresés", command=search_ceg)
search_button.grid(row=0, column=2, sticky="w", padx=10, pady=10,)

default_button = tk.Button(data_frame, font=("Arial", 16), text="Keresés alaphelyzetbe állítása", command=default_search)
default_button.grid(row=0, column=4, sticky="nsew", padx=10, pady=10,)

empty_button = tk.Button(data_frame, font=("Arial", 16), text="Beviteli mezők kiürítése", command=empty_entry)
empty_button.grid(row=2, column=4, sticky="nsew", padx=10, pady=10,)

# Események kezelése
treeview.bind("<ButtonRelease-1>", show_selected_column)  # Bal egérkattintásra oszlop kijelzés
treeview.bind("<Double-1>", lambda event: copy_to_clipboard())  # Dupla kattintás a másoláshoz
treeview.bind("<<TreeviewSelect>>", on_tree_select)  # Kijelölt sor után QR kód generálás
search_entry.bind("<Return>", lambda event: search_ceg())
root.bind("<Control-Delete>", lambda event: del_button.invoke())

# QR kód megjelenítő canvas
qr_canvas = tk.Canvas(qr_frame, bg="white", width=250, height=250)
qr_canvas.grid(row=0, column=0, padx=10, pady=10, sticky="nsew")

# QR kód szöveg
qr_text = tk.StringVar()
qr_text.set("Válassz egy sort!")
qr_text_label = tk.Label(qr_frame, textvariable=qr_text, bg="white", wraplength=700, justify="center", font=("Arial", 16))
qr_text_label.grid(row=3, column=0, pady=10, sticky="nsew")

# Fő ciklus
root.mainloop()